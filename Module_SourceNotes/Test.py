import pdfplumber
from docx import Document

import os
from google import genai
from dotenv import load_dotenv



#I only created this so I can test how the pdf reader works

#file:///

url = r"https://learn.madisoncollege.edu/d2l/le/dropbox/77056/202659/DownloadAttachment?fid=6718639"




word_doc = r"C:\Users\basil\Downloads\Botany of Baking Updated Sp 26.docx"

def pdf_transcript_v2(user_input):
    all_pages = []

    if "file:///" in user_input:
        user_input = user_input.replace("file:///", "")
    
    file_name = user_input.rsplit('\\')[-1].rsplit('.')[0]

    
    with pdfplumber.open(user_input) as docx:
         for page in docx.pages:
            formatted = f"\n\n-- Page: {page.page_number} --\n"
            all_pages.append(formatted)
            all_pages.append(page.extract_text())
            

         full_text = "\n".join(all_pages)
         print(full_text)
         


def word_docx_transcript(user_input):
    document = Document(user_input)
    all_pages = []

    file_name = user_input.rsplit("\\")[-1].rsplit(".")[0]

    for paragraph in document.paragraphs:
        print(paragraph.text)
        all_pages.append(paragraph.text)
        
    
    for table in document.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                all_pages.append(row_text)
    
    full_text = "\n".join(all_pages)


file_path = r"C:\Users\basil\Downloads\Chris Peterson_CompSci2_Fall 25.pdf"

load_dotenv()


def gemini_extract(user_input):
    try:
        api_key = os.getenv("GEMINI_API_KEY")

        client = genai.Client(api_key=api_key)

        uploaded_file = client.files.upload(file=user_input)


        prompt = """
        Extract the exact contents of this file, include no extra commentary just organized text from the following file.
        """

        response = client.models.generate_content(
        model="gemini-2.5-flash-lite",
        contents=[uploaded_file, prompt]
    )

        return response.text

    except Exception as e:
        print(f"Gemini Extraction failed: {e}")
        return None


print(gemini_extract(file_path))








