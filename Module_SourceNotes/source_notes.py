# import statements
from youtube_transcript_api import YouTubeTranscriptApi
from Module_SourceNotes.source_note_data import SourceNote
from datetime import datetime
import uuid
import yt_dlp #pip install yt-dlp    <----- import statement
import re
import shutil #To save the files in a specific folder
import fitz        #run pip install pymupdf requests pytesseract pdf2image pillow to install ibraries
import requests
import pytesseract
from pdf2image import convert_from_path, convert_from_bytes
from PIL import Image
import os
import io
from urllib.parse import unquote
from urllib.request import url2pathname 

import pdfplumber
from docx import Document

import os
from google import genai
from dotenv import load_dotenv
load_dotenv()


pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

class SourceNotes_Extractor:
    def __init__(self):
        self.source_path = ""
        self.filename = ""
        self.completed_transcript = ""



    def data_instance(self, current_source_type): # 
        return SourceNote(
            id = str(uuid.uuid4()),
            title = self.filename,
            source_type = current_source_type,
            source_path = self.source_path,
            transcript = self.completed_transcript,
            created_at = datetime.now()
        )
    
    def save_transcript(self):
        try:
            shutil.move(self.filename,"Module_SourceNotes\\Exported_Transcripts") # takes in 2 parameters, the name of the file, and where you are moving it to
            print(self.filename)
            print(f"And was able to move {self.filename} to its destination")
        except FileNotFoundError:
            print(f"Error: The source file '{self.filename}' was not found")
        except Exception as e:
            print(f"Found error: {e}")



    def extract_transcript(self, user_input): #logic may need fixing aswell as in main needs exception handling
        try:

            if "youtube.com" in user_input or "youtu.be" in user_input:
                return self.youtube_transcript(user_input)
            
            elif ".pdf" in user_input:
                current_source_note = self.pdf_transcript_v2(user_input)
                if not current_source_note or len(current_source_note.transcript) < 100:
                    return self.gemini_extract(user_input)
                
                else:
                    return current_source_note
            
            elif user_input.endswith(".docx"):
                current_source_note = self.word_docx_transcript(user_input)
                if not current_source_note or len(current_source_note.transcript) < 100:
                    return self.gemini_extract(user_input)
                
                else:
                    return current_source_note
            
            else:
                return None
        
        except Exception as e:
            print(f"Source note extraction failed: {e}")
            return None

# -----------------------------------------MANUAL_INPUT_CODE-----------------------------------------

    def manual_text_transcript(self, pasted_user_text, pasted_user_title):
        self.completed_transcript = pasted_user_text
        self.filename = pasted_user_title
        self.source_path = "None"

        return self.data_instance("Raw Text")
        


# -----------------------------------------YOUTUBE_TRANSCRIPTS_CODE-----------------------------------------
    def youtube_vid_title(self):
        ydl_opts = {'quiet': True} #suppresses console output from yt-dlp while downloading the vedio

        with yt_dlp.YoutubeDL(ydl_opts) as ydl: 
            info = ydl.extract_info(self.source_path, download=False) #fetches vedios metadata without having to download
            yt_title = info.get('title')

        yt_title = re.sub(r'[<>:"/\\|?*\x00-\x1F]+', "_", yt_title)#replaces illegal sybmols with _
        return yt_title.strip().strip(".")


    def youtube_transcript(self, user_input):
        self.source_path = user_input
        if "v=" in self.source_path:
            video_id = self.source_path.split("v=")[1].split("&")[0] #Takes the Yt url and splits the list into two halves, where everything before v= is 0 and after is 1 and then it's repeated.
        else:
            video_id = self.source_path.split("/")[-1] #does the same thing splitting the link everywhere with a '/' and then grabbing the last element which is the id
        
        youtube = YouTubeTranscriptApi()
        transcript = youtube.fetch(video_id) #returns a list of transcript segments objects: .text, .start, .duration

        # Join each transcript segment on a new line instead of a single long line
        completed_transcript = "\n".join(
            f"{item.start:.2f}s | {item.duration:.2f}s | {item.text}"
            for item in transcript
        )

        title = self.youtube_vid_title() or video_id

        self.filename = f"{title}-Youtube_transcript.txt" #custom file name

        self.completed_transcript = completed_transcript
        return self.data_instance("youtube")

    

# -----------------------------------------PDF_TRANSCRIPTS_CODE-----------------------------------------
        """
        Library                           What it does
        ---------------------------------------------------------------
        pymupdf (imported as fitz)       Reads PDFs, extracts text per page
        requests                         Downloads PDFs from a URL
        pytesseract                      OCR engine — reads text from images
        pdf2image                        Converts PDF pages into images (for OCR)
        pillow (PIL)                     Image processing, required by pytesseract
        """


    def pdf_transcript(self, pdf_input=""): #get the pdf into the memory
        
        if pdf_input.startswith("http://") or pdf_input.startswith("https://"):
            pdf_bytes = self._download_pdf(pdf_input)   # fetch and download pdf from url and returns raw bytes
            pdf_source = "url"
        elif pdf_input.startswith("file:///"):                       
            local_path = url2pathname(pdf_input[7:])                  # strips "file://" and decodes %20 etc.
            pdf_input = local_path
            pdf_bytes = None
            pdf_source = "local"
        else:           #if pdf isn't a url sets the bytes to none assumes the source is local path 
            pdf_bytes = None
            pdf_source = "local"

        #STEP2: open pdf with pymupdf
        if pdf_source == "url":
            doc = fitz.open(stream=pdf_bytes, filetype="pdf") #tells PyMuPDF to open the PDF directly from those bytes in memory without needing a file on disk
                #stream=pdf_bytes = passes the raw bytes
                #filetype="pdf" =  tells fitz to interpret those bytes as a pdf
        else:
            doc = fitz.open(pdf_input) #directly open the file if its local and not a url

        #STEP3: extract the pdf page by page
        all_pages = []
        for page_num in range(len(doc)):  # loop through every page 
            page = doc[page_num]
            text = page.get_text()  # attempt direct text extraction

            if self._is_empty(text):
                print(f"page{page_num+1} has been scanned, running OCR...")
                text = self._ocr_page(page, pdf_bytes, pdf_input, page_num, pdf_source)

                # Format each page clearly
            formatted = f"Page {page_num+1} --\n{text.strip()}\n"
            all_pages.append(formatted)
        doc.close()

        #STEP4: save the extracted text 
        full_text = "\n".join(all_pages)
        self.filename = self._get_pdf_filename(pdf_input)
        self.source_path = pdf_input          # add this
        self.completed_transcript = full_text
        
        
        self.save_transcript()
        return self.data_instance("pdf")
    
    def _download_pdf(self, url): #Downloads a PDF from a URL and returns it as raw bytes
        response = requests.get(url) #"request" is alibrary to access the internet this gets the url and stores it in response
        response.raise_for_status() # cheks if the download worked and raises an error if download failed
        return response.content #.content = raw bites of pdf stores pdf in memory
    
    def _is_empty(self,text): #checks if the page is empty 
        return not text.strip()#After stripping, if the result is an empty string "" python treats that as False not False = True  So the method returns True when the text is empty and false when it has real content this is how the main loop knows to run OCR or not
        
    def _ocr_page(self, page, pdf_bytes, pdf_input, page_num, pdf_source):#Converts a single PDF page to an image and runs OCR on it to read the pdf
        if pdf_source == "url":
            images = convert_from_bytes(pdf_bytes, first_page=page_num+1, last_page=page_num+1)
        else:
            images = convert_from_path(pdf_input, first_page=page_num+1, last_page=page_num+1)

        return pytesseract.image_to_string(images[0])#It's essentially reading the image like a human would read a page
        #images[0] convert_from_* always returns a list of images since we only asked for one page, we grab the first and only item with [0]
    
    def _get_pdf_filename(self, pdf_input):#This method creates a clean, readable filename for the output text file.
        if pdf_input.startswith("http"):
            # Safer fix:
            base = pdf_input.split("/")[-1].split("?")[0].replace(".pdf", "")#.split("/") breaks the URL into a list wherever there's a /
            #[-1] grabs the last item in that list which is the filename "invoice.pdf" .replace(".pdf", "") removes the .pdf extension, leaving just "invoice"
        else:   
            base = os.path.basename(pdf_input).replace(".pdf", "")
        return f"{base}-PDF-transcript.txt"



# -----------------------------------------PDF_TRANSCRIPTS_CODE - V2-----------------------------------------
    def pdf_transcript_v2(self, user_input):
        all_pages = []

        if "file:///" in user_input:
            user_input = user_input.replace("file:///", "")
        
        self.filename = user_input.rsplit('\\')[-1].rsplit('.')[0]

        
        with pdfplumber.open(user_input) as docx:
            for page in docx.pages:
                formatted = f"\n\n-- Page: {page.page_number} --\n"
                all_pages.append(formatted)
                all_pages.append(page.extract_text())
                

            self.completed_transcript = "\n".join(all_pages)
            self.source_path = user_input

            return self.data_instance(".pdf")



# # -----------------------------------------DOCX_TRANSCRIPTS_CODE-----------------------------------------
    def word_docx_transcript(self, user_input):
        document = Document(user_input)
        all_pages = []

        self.filename = user_input.rsplit("\\")[-1].rsplit(".")[0]

        for paragraph in document.paragraphs:
            all_pages.append(paragraph.text)
            
        
        for table in document.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    all_pages.append(row_text)
        
        self.completed_transcript = "\n".join(all_pages)
        self.source_path = user_input

        return self.data_instance(".docx")



# # -----------------------------------------GEMINI_TRANSCRIPTS_CODE-----------------------------------------
    def gemini_extract(self, user_input):
        try:
            api_key = os.getenv("GEMINI_API_KEY")

            client = genai.Client(api_key=api_key)

            uploaded_file = client.files.upload(file=user_input)


            prompt = """
            Extract the exact contents of this file, include no extra commentary just organized text from the following file.
            """

            response = client.models.generate_content(
            model="gemini-2.5-flash-lite",
            contents=[uploaded_file, prompt]
        )

            self.completed_transcript = response.text
            self.filename = user_input.rsplit('\\')[-1].rsplit('.')[0]
            self.source_path = user_input

            return self.data_instance("Gemini_extraction")

        except Exception as e:
            print(f"Gemini Extraction failed: {e}")
            return None
            

